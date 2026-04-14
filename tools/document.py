from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Union

from office_agent.utils import ensure_parent_dir

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from openpyxl import Workbook, load_workbook
except ImportError:  # pragma: no cover
    Workbook = None
    load_workbook = None


def read_word_document(path: str) -> str:
    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，请先安装 requirements.txt。")
    doc = Document(path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())


def create_word_document(content: str, path: str, title: str | None = None) -> bool:
    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，请先安装 requirements.txt。")
    target = ensure_parent_dir(path)
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for block in content.split("\n"):
        doc.add_paragraph(block)
    doc.save(target)
    return True


def append_to_word(path: str, content: str) -> bool:
    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，请先安装 requirements.txt。")
    doc = Document(path)
    for block in content.split("\n"):
        doc.add_paragraph(block)
    doc.save(path)
    return True


def read_excel_document(path: str, sheet: Union[str, int] = 0) -> Dict[str, Any]:
    if load_workbook is None:
        raise RuntimeError("缺少 openpyxl 依赖，请先安装 requirements.txt。")
    workbook = load_workbook(path)
    worksheet = workbook[sheet] if isinstance(sheet, str) else workbook.worksheets[sheet]
    rows = list(worksheet.iter_rows(values_only=True))
    headers = [str(cell) if cell is not None else "" for cell in rows[0]] if rows else []
    data = [list(row) for row in rows[1:]] if len(rows) > 1 else []
    return {
        "headers": headers,
        "data": data,
        "sheet_names": workbook.sheetnames,
        "active_sheet": worksheet.title,
    }


def create_excel_document(data: List[List[Any]], path: str, headers: List[str] | None = None) -> bool:
    if Workbook is None:
        raise RuntimeError("缺少 openpyxl 依赖，请先安装 requirements.txt。")
    target = ensure_parent_dir(path)
    workbook = Workbook()
    worksheet = workbook.active
    if headers:
        worksheet.append(headers)
    for row in data:
        worksheet.append(row)
    workbook.save(target)
    return True


def update_excel_cell(path: str, sheet: str, row: int, col: int, value: Any) -> bool:
    if load_workbook is None:
        raise RuntimeError("缺少 openpyxl 依赖，请先安装 requirements.txt。")
    workbook = load_workbook(path)
    worksheet = workbook[sheet]
    worksheet.cell(row=row, column=col, value=value)
    workbook.save(path)
    return True
